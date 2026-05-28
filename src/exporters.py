"""Bộ chuyển đổi Markdown → HTML / LaTeX / DOCX / PDF."""
import os
import re
import html


# ============================================================
# HTML  (re-export từ preview để main_window dùng một import duy nhất)
# ============================================================
from .preview import markdown_to_html  # noqa: F401  (re-export)
from .i18n import t
from .text_normalizer import normalize_markdown_for_compile


# ============================================================
# LaTeX --- eMeX modern template
# ============================================================
# Thiết kế: chữ Latin Modern + microtype, bảng màu thương hiệu eMeX
# (xanh #2563EB / tím #7C3AED), heading số có màu, bảng booktabs,
# trích dẫn dạng callout (tcolorbox), code block có thanh màu trái và
# tô ngôn ngữ, header/footer tối giản. Có sẵn theorem env tiếng Việt.
LATEX_TEMPLATE = r"""%==============================================================================
% eMeX --- Modern LaTeX Template
% Tài liệu được sinh tự động từ Markdown bằng eMeX.
% Dự án: https://github.com/nhhai-math/eMeX
% Gợi ý biên dịch: lualatex / xelatex cho tiếng Việt tốt nhất;
% pdflatex vẫn dùng được nhờ babel-vietnamese.
%==============================================================================
\documentclass[11pt,a4paper]{article}

%--- Encoding & ngôn ngữ ---
\usepackage[utf8]{inputenc}
\usepackage[T5,T1]{fontenc}
\usepackage[english,vietnamese]{babel}

%--- Typography ---
\usepackage{lmodern}
\usepackage{microtype}
\usepackage{parskip}

%--- Bố cục trang ---
\usepackage[a4paper,margin=2.4cm,headheight=15pt,footskip=1.2cm]{geometry}

%--- Toán học ---
\usepackage{amsmath,amssymb,amsthm,mathtools}
\usepackage{bm}

%--- Đồ hoạ & màu ---
\usepackage{graphicx}
\usepackage[dvipsnames,table]{xcolor}

%--- Bảng màu thương hiệu eMeX ---
\definecolor{emexPrimary}{HTML}{2563EB}
\definecolor{emexAccent}{HTML}{7C3AED}
\definecolor{emexInk}{HTML}{0F172A}
\definecolor{emexBody}{HTML}{1F2937}
\definecolor{emexSlate}{HTML}{475569}
\definecolor{emexMute}{HTML}{64748B}
\definecolor{emexLine}{HTML}{E2E8F0}
\definecolor{emexBg}{HTML}{F8FAFC}
\definecolor{emexCode}{HTML}{F1F5F9}
\definecolor{emexCodeKw}{HTML}{7C3AED}
\definecolor{emexCodeStr}{HTML}{059669}
\definecolor{emexCodeCmt}{HTML}{64748B}
\color{emexBody}

%--- TikZ ---
\usepackage{tikz}
\usetikzlibrary{arrows.meta,positioning,calc,decorations.pathreplacing,shapes.geometric}

%--- Danh sách ---
\usepackage{enumitem}
\setlist[itemize]{leftmargin=1.5em,itemsep=0.25em,topsep=0.4em,parsep=0pt}
\setlist[enumerate]{leftmargin=1.7em,itemsep=0.25em,topsep=0.4em,parsep=0pt}

%--- Bảng ---
\usepackage{booktabs}
\usepackage{array}
\usepackage{tabularx}
\renewcommand{\arraystretch}{1.3}
\arrayrulecolor{emexLine}

%--- Mã nguồn ---
\usepackage{listings}
\lstdefinelanguage{JavaScript}{
  morekeywords={typeof, new, true, false, catch, function, return, null, switch,
    var, let, const, if, in, while, do, else, case, break, async, await, of,
    for, class, extends, import, export, from, this, super, default, try},
  sensitive=true,
  morecomment=[l]{//},
  morecomment=[s]{/*}{*/},
  morestring=[b]',
  morestring=[b]",
  morestring=[b]`,
}
\lstdefinestyle{emexcode}{
  basicstyle=\ttfamily\footnotesize\color{emexBody},
  backgroundcolor=\color{emexCode},
  frame=leftline,
  framesep=10pt,
  framerule=2.5pt,
  rulecolor=\color{emexPrimary},
  xleftmargin=0pt,
  breaklines=true,
  breakatwhitespace=false,
  breakindent=0pt,
  columns=fullflexible,
  keepspaces=true,
  keywordstyle=\bfseries\color{emexCodeKw},
  ndkeywordstyle=\color{emexCodeKw},
  commentstyle=\itshape\color{emexCodeCmt},
  stringstyle=\color{emexCodeStr},
  numbers=none,
  showstringspaces=false,
  tabsize=2,
  aboveskip=10pt,
  belowskip=10pt,
}
\lstset{style=emexcode}

%--- Strike-through ---
\usepackage[normalem]{ulem}

%--- Hyperref (giữ ở cuối preamble, trước tcolorbox) ---
\usepackage[
  colorlinks=true,
  linkcolor=emexPrimary,
  urlcolor=emexPrimary,
  citecolor=emexAccent,
  pdfborder={0 0 0},
  bookmarksopen=true,
  bookmarksnumbered=true,
  pdfusetitle,
  pdfcreator={eMeX --- Markdown editor for mathematics},
]{hyperref}

%--- Header / footer ---
\usepackage{fancyhdr}
\pagestyle{fancy}
\fancyhf{}
\renewcommand{\headrulewidth}{0pt}
\renewcommand{\footrulewidth}{0pt}
\fancyhead[R]{\small\color{emexMute}\itshape\nouppercase{\leftmark}}
\fancyfoot[L]{\footnotesize\color{emexMute}\textsc{eMeX}}
\fancyfoot[R]{\footnotesize\color{emexMute}--~\thepage~--}

%--- Heading styles ---
\usepackage{titlesec}
\titleformat{\section}
  {\normalfont\Large\bfseries\color{emexInk}}
  {\textcolor{emexPrimary}{\thesection}}{0.8em}{}
\titlespacing*{\section}{0pt}{1.6em}{0.6em}

\titleformat{\subsection}
  {\normalfont\large\bfseries\color{emexInk}}
  {\textcolor{emexPrimary}{\thesubsection}}{0.7em}{}
\titlespacing*{\subsection}{0pt}{1.1em}{0.45em}

\titleformat{\subsubsection}
  {\normalfont\normalsize\bfseries\color{emexSlate}}
  {\textcolor{emexPrimary}{\thesubsubsection}}{0.6em}{}
\titlespacing*{\subsubsection}{0pt}{0.9em}{0.35em}

\titleformat{\paragraph}[runin]
  {\normalfont\normalsize\bfseries\color{emexAccent}}{}{0pt}{}[\,]
\titleformat{\subparagraph}[runin]
  {\normalfont\normalsize\itshape\color{emexSlate}}{}{0pt}{}[\,]

%--- Callout / quote ---
\usepackage[most]{tcolorbox}
\newtcolorbox{emexquote}{
  enhanced jigsaw,
  colback=emexBg,
  colframe=emexPrimary,
  boxrule=0pt,
  leftrule=3pt,
  arc=2pt,
  outer arc=2pt,
  left=12pt,right=12pt,top=8pt,bottom=8pt,
  fontupper=\itshape\color{emexSlate},
  before skip=10pt, after skip=10pt,
}

%--- Môi trường định lý (sẵn có để dùng) ---
\theoremstyle{plain}
\newtheorem{theorem}{Định lý}[section]
\newtheorem{proposition}[theorem]{Mệnh đề}
\newtheorem{lemma}[theorem]{Bổ đề}
\newtheorem{corollary}[theorem]{Hệ quả}
\theoremstyle{definition}
\newtheorem{definition}[theorem]{Định nghĩa}
\newtheorem{example}[theorem]{Ví dụ}
\newtheorem{exercise}[theorem]{Bài tập}
\theoremstyle{remark}
\newtheorem*{remark}{Chú ý}

%--- Inline code helper ---
\newcommand{\emexinlinecode}[1]{{\ttfamily\small\color{emexAccent}#1}}

%--- Task list markers ---
\newcommand{\emextaskdone}{\textcolor{emexPrimary}{$\boxtimes$}}
\newcommand{\emextasktodo}{\textcolor{emexMute}{$\square$}}

%--- Tiêu đề tài liệu ---
\makeatletter
\renewcommand{\maketitle}{%
  \begin{flushleft}
    \vspace*{-0.6cm}
    {\color{emexPrimary}\rule{4em}{2.5pt}}\par
    \vspace{8pt}
    {\Huge\bfseries\color{emexInk}\@title\par}
    \vspace{6pt}
    {\normalsize\color{emexSlate}\@author~~\textbullet~~\@date}\par
    \vspace{8pt}
    {\color{emexLine}\rule{\linewidth}{0.4pt}}\par
    \vspace{14pt}
  \end{flushleft}
}
\makeatother

\title{@@TITLE@@}
\author{eMeX}
\date{\today}

\begin{document}
\maketitle

@@BODY@@

\end{document}
"""


# Mapping ngôn ngữ code fence → tên hợp lệ trong listings
_LST_LANG_MAP = {
    'python': 'Python', 'py': 'Python',
    'c': 'C', 'cpp': 'C++', 'c++': 'C++',
    'csharp': '[Sharp]C', 'cs': '[Sharp]C', 'c#': '[Sharp]C',
    'java': 'Java',
    'js': 'JavaScript', 'javascript': 'JavaScript',
    'ts': 'JavaScript', 'typescript': 'JavaScript',
    'html': 'HTML', 'xml': 'XML', 'css': 'HTML',
    'php': 'PHP',
    'sql': 'SQL',
    'bash': 'bash', 'sh': 'bash', 'shell': 'bash', 'zsh': 'bash',
    'tex': '[LaTeX]TeX', 'latex': '[LaTeX]TeX',
    'r': 'R',
    'matlab': 'Matlab',
    'ruby': 'Ruby', 'rb': 'Ruby',
    'perl': 'Perl',
    'mathematica': 'Mathematica',
}


def _latex_escape(text):
    """Escape ký tự đặc biệt LaTeX trong văn bản thường."""
    if not text:
        return text
    # backslash phải xử lý trước
    text = text.replace('\\', '\x00BS\x00')
    pairs = [
        ('&', r'\&'),
        ('%', r'\%'),
        ('$', r'\$'),
        ('#', r'\#'),
        ('_', r'\_'),
        ('{', r'\{'),
        ('}', r'\}'),
        ('~', r'\textasciitilde{}'),
        ('^', r'\textasciicircum{}'),
    ]
    for o, n in pairs:
        text = text.replace(o, n)
    text = text.replace('\x00BS\x00', r'\textbackslash{}')
    return text


def _latex_inline(text):
    """Convert markdown inline (sau khi đã escape) thành LaTeX."""
    # **bold** / __bold__
    text = re.sub(r'\*\*([^\*\n]+?)\*\*', r'\\textbf{\1}', text)
    text = re.sub(r'__([^_\n]+?)__', r'\\textbf{\1}', text)
    # *italic* / _italic_ (không khớp ** vì đã xử lý ở trên)
    text = re.sub(r'(?<![\*\\])\*([^\*\n]+?)\*(?!\*)', r'\\textit{\1}', text)
    text = re.sub(r'(?<![_a-zA-Z\\])_([^_\n]+?)_(?![_a-zA-Z])', r'\\textit{\1}', text)
    # ~~strike~~
    text = re.sub(r'~~([^~\n]+?)~~', r'\\sout{\1}', text)
    # [text](url) – chú ý đã escape '_' trong url, cần unescape một số
    def link_repl(m):
        label = m.group(1)
        url = (m.group(2) or m.group(3)).replace(r'\_', '_').replace(r'\#', '#').replace(r'\%', '%')
        return f"\\href{{{url}}}{{{label}}}"
    text = re.sub(r'\[([^\]]+)\]\((?:<([^>]*)>|([^)\s]+))\)', link_repl, text)
    return text


def _strip_leading_title_h1(source, title):
    """Bỏ H1 đầu tiên nếu trùng tiêu đề tài liệu (tránh hiện 2 lần)."""
    if not title:
        return source
    title_norm = title.strip()
    lines = source.split('\n')
    for idx, line in enumerate(lines):
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith('# ') and stripped[2:].strip() == title_norm:
            del lines[idx]
            if idx < len(lines) and not lines[idx].strip():
                del lines[idx]
            return '\n'.join(lines)
        return source
    return source


def _markdown_to_latex_body(source):
    """Convert raw markdown body (chưa wrap document) → LaTeX body."""
    placeholders = []

    def stash(kind, content):
        placeholders.append((kind, content))
        return f"@@PH{len(placeholders) - 1}@@"

    # 1. TikZ block (ưu tiên trước code thường)
    def tikz_repl(m):
        return "\n" + stash('tikz', m.group(1).strip()) + "\n"
    source = re.sub(r'```[ \t]*tikz[ \t]*\r?\n([\s\S]*?)```',
                    tikz_repl, source, flags=re.IGNORECASE)

    # 2. Code fences (cho phép C++, C# trong tên ngôn ngữ)
    def code_repl(m):
        return "\n" + stash('code', (m.group(1) or '', m.group(2))) + "\n"
    source = re.sub(r'```([\w+#\-]*)[^\S\r\n]*\r?\n([\s\S]*?)```',
                    code_repl, source)

    # 3. Display math $$...$$
    def dmath_repl(m):
        return "\n" + stash('dmath', m.group(1).strip()) + "\n"
    source = re.sub(r'\$\$([\s\S]+?)\$\$', dmath_repl, source)

    # 4. Inline math $...$
    def imath_repl(m):
        return stash('imath', m.group(1))
    source = re.sub(r'(?<!\$)\$([^\n$]+?)\$(?!\$)', imath_repl, source)

    # 5. Inline code
    def icode_repl(m):
        return stash('icode', m.group(1))
    source = re.sub(r'`([^`\n]+)`', icode_repl, source)

    # ----- Tokenize line-by-line -----
    lines = source.split('\n')
    out = []
    i = 0
    in_list = None  # 'itemize' | 'enumerate' | None

    def close_list():
        nonlocal in_list
        if in_list:
            out.append(f"\\end{{{in_list}}}")
            in_list = None

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Dòng có chứa duy nhất placeholder dạng @@PHn@@ → in raw
        if re.fullmatch(r'@@PH\d+@@', stripped):
            close_list()
            out.append(stripped)
            i += 1
            continue

        # Heading
        h = re.match(r'^(#{1,6})\s+(.*)$', stripped)
        if h:
            close_list()
            level = len(h.group(1))
            content = _latex_inline(_latex_escape(h.group(2)))
            cmd = ['section', 'subsection', 'subsubsection',
                   'paragraph', 'subparagraph', 'subparagraph'][min(level - 1, 5)]
            out.append(f"\\{cmd}{{{content}}}")
            i += 1
            continue

        # Horizontal rule
        if re.match(r'^(?:---+|\*\*\*+|___+)\s*$', stripped):
            close_list()
            out.append(r"\par\vspace{0.6em}\noindent{\color{emexLine}\rule{\linewidth}{0.6pt}}\par\vspace{0.6em}")
            i += 1
            continue

        # Task list: - [ ] / - [x]
        task = re.match(r'^[-*+]\s+\[([ xX])\]\s+(.*)$', stripped)
        if task:
            if in_list != 'itemize':
                close_list()
                out.append(r"\begin{itemize}")
                in_list = 'itemize'
            mark = r"\emextaskdone" if task.group(1).lower() == 'x' else r"\emextasktodo"
            content = _latex_inline(_latex_escape(task.group(2)))
            out.append(f"  \\item[{mark}] {content}")
            i += 1
            continue

        # Bullet list
        bullet = re.match(r'^[-*+]\s+(.*)$', stripped)
        if bullet:
            if in_list != 'itemize':
                close_list()
                out.append(r"\begin{itemize}")
                in_list = 'itemize'
            content = _latex_inline(_latex_escape(bullet.group(1)))
            out.append(f"  \\item {content}")
            i += 1
            continue

        # Ordered list
        om = re.match(r'^\d+\.\s+(.*)$', stripped)
        if om:
            if in_list != 'enumerate':
                close_list()
                out.append(r"\begin{enumerate}")
                in_list = 'enumerate'
            content = _latex_inline(_latex_escape(om.group(1)))
            out.append(f"  \\item {content}")
            i += 1
            continue

        # Block quote → gộp các dòng liên tiếp (cùng đoạn) bằng emexquote
        if stripped.startswith('>'):
            close_list()
            paragraphs = []
            current = []
            while i < len(lines):
                ls = lines[i].strip()
                if not ls.startswith('>'):
                    break
                quote_text = ls[1:]
                if quote_text.startswith(' '):
                    quote_text = quote_text[1:]
                if not quote_text:
                    if current:
                        paragraphs.append(' '.join(current))
                        current = []
                    i += 1
                    continue
                current.append(quote_text)
                i += 1
            if current:
                paragraphs.append(' '.join(current))
            if not paragraphs:
                out.append("")
                continue
            out.append(r"\begin{emexquote}")
            for idx, p in enumerate(paragraphs):
                if idx > 0:
                    out.append('')
                out.append(_latex_inline(_latex_escape(p)))
            out.append(r"\end{emexquote}")
            continue

        # GFM Table với booktabs + alignment
        if (stripped.startswith('|') and i + 1 < len(lines)
                and re.match(r'^\|[\s\-:|]+\|\s*$', lines[i + 1].strip())):
            close_list()
            header = [c.strip() for c in stripped.strip().strip('|').split('|')]
            sep_cells = [c.strip() for c in lines[i + 1].strip().strip('|').split('|')]
            align = []
            for c in sep_cells:
                left_colon = c.startswith(':')
                right_colon = c.endswith(':')
                if left_colon and right_colon:
                    align.append('c')
                elif right_colon:
                    align.append('r')
                else:
                    align.append('l')
            i += 2
            rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                cells = [c.strip() for c in lines[i].strip().strip('|').split('|')]
                rows.append(cells)
                i += 1
            n = len(header)
            align = (align + ['l'] * n)[:n]
            spec = ''.join(align)
            out.append(r"\begin{center}")
            out.append(f"\\begin{{tabular}}{{{spec}}}")
            out.append(r"\toprule")
            header_cells = ' & '.join(
                f"\\textbf{{\\color{{emexInk}}{_latex_inline(_latex_escape(c))}}}"
                for c in header)
            out.append(header_cells + r' \\')
            out.append(r'\midrule')
            for row in rows:
                row = (row + [''] * n)[:n]
                out.append(' & '.join(_latex_inline(_latex_escape(c)) for c in row) + r' \\')
            out.append(r'\bottomrule')
            out.append(r"\end{tabular}")
            out.append(r"\end{center}")
            out.append("")
            continue

        # Image standalone: ![alt](url)
        img_only = re.match(r'^!\[([^\]]*)\]\((?:<([^>]*)>|([^)\s]+))\)\s*$', stripped)
        if img_only:
            close_list()
            alt_raw = img_only.group(1)
            alt = _latex_inline(_latex_escape(alt_raw)) if alt_raw else ''
            url = img_only.group(2) or img_only.group(3)
            if url.startswith('http://') or url.startswith('https://'):
                # URL ngoài: .tex độc lập không tải được, chỉ in link
                out.append(r"\begin{center}")
                if alt:
                    out.append(f"\\textit{{{alt}}}\\\\")
                out.append(f"\\href{{{url}}}{{\\nolinkurl{{{url}}}}}")
                out.append(r"\end{center}")
            else:
                out.append(r"\begin{figure}[ht]")
                out.append(r"  \centering")
                out.append(f"  \\includegraphics[width=0.7\\linewidth]{{{url}}}")
                if alt:
                    out.append(f"  \\caption{{{alt}}}")
                out.append(r"\end{figure}")
            i += 1
            continue

        # Dòng trống → đóng list
        if not stripped:
            close_list()
            out.append("")
            i += 1
            continue

        # Đoạn văn thường: gom các dòng liên tiếp không phải block đặc biệt
        para = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt:
                break
            # Nếu dòng kế là heading / list / quote / table thì dừng
            if (re.match(r'^(#{1,6})\s', nxt) or re.match(r'^[-*+]\s', nxt)
                    or re.match(r'^\d+\.\s', nxt) or nxt.startswith('>')
                    or nxt.startswith('|') or nxt.startswith('```')
                    or re.fullmatch(r'@@PH\d+@@', nxt)):
                break
            para.append(nxt)
            i += 1
        close_list()
        joined = ' '.join(para)
        out.append(_latex_inline(_latex_escape(joined)))
        out.append("")

    close_list()
    body = '\n'.join(out)

    # ----- Restore placeholders -----
    def restore_tikz(content):
        content = content.strip()
        if re.search(r'\\begin\s*\{\s*tikzpicture\s*\}', content):
            return f"\n\\begin{{center}}\n{content}\n\\end{{center}}\n"
        return (f"\n\\begin{{center}}\n\\begin{{tikzpicture}}\n"
                f"{content}\n\\end{{tikzpicture}}\n\\end{{center}}\n")

    def restore(m):
        idx = int(m.group(1))
        kind, content = placeholders[idx]
        if kind == 'tikz':
            return restore_tikz(content)
        if kind == 'code':
            lang_raw, code = content
            lang_key = (lang_raw or '').lower().strip()
            lst_lang = _LST_LANG_MAP.get(lang_key)
            if lst_lang:
                return (f"\n\\begin{{lstlisting}}[language={lst_lang}]\n"
                        f"{code.rstrip()}\n\\end{{lstlisting}}\n")
            return f"\n\\begin{{lstlisting}}\n{code.rstrip()}\n\\end{{lstlisting}}\n"
        if kind == 'dmath':
            return f"\n\\[\n{content}\n\\]\n"
        if kind == 'imath':
            return f"${content}$"
        if kind == 'icode':
            return f"\\emexinlinecode{{{_latex_escape(content)}}}"
        return ''
    body = re.sub(r'@@PH(\d+)@@', restore, body)
    return body


def markdown_to_latex(source, title="Tài liệu Markdown"):
    """Convert markdown source → standalone LaTeX document."""
    source = normalize_markdown_for_compile(source)
    source = _strip_leading_title_h1(source, title)
    body = _markdown_to_latex_body(source)
    return (LATEX_TEMPLATE
            .replace('@@TITLE@@', _latex_escape(title))
            .replace('@@BODY@@', body))


# ============================================================
# DOCX
# ============================================================
def markdown_to_docx(source, path):
    """Convert markdown → .docx (cần python-docx)."""
    source = normalize_markdown_for_compile(source)
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    except ImportError as exc:
        raise RuntimeError(
            t("Chưa có thư viện python-docx. Cài bằng:\n  pip install python-docx")) from exc

    doc = Document()

    # Default font
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)

    lines = source.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Display math: $$ ... $$ -> Word OMML XML.
        if stripped.startswith('$$'):
            math_lines = []
            first = stripped[2:]
            if first.endswith('$$') and len(first) >= 2:
                math_lines.append(first[:-2])
                i += 1
            else:
                if first:
                    math_lines.append(first)
                i += 1
                while i < len(lines):
                    candidate = lines[i].strip()
                    if candidate.endswith('$$'):
                        math_lines.append(candidate[:-2])
                        i += 1
                        break
                    math_lines.append(lines[i])
                    i += 1
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            _docx_append_math(p, '\n'.join(math_lines).strip())
            continue

        # Heading
        h = re.match(r'^(#{1,6})\s+(.*)$', stripped)
        if h:
            level = len(h.group(1))
            p = doc.add_heading(level=min(level, 9))
            _docx_fill_runs(p, h.group(2))
            i += 1
            continue

        # Code block
        if stripped.startswith('```'):
            lang = stripped[3:].strip()
            i += 1
            buf = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                buf.append(lines[i])
                i += 1
            i += 1  # skip closing
            p = doc.add_paragraph()
            run = p.add_run('\n'.join(buf))
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            if lang:
                p.paragraph_format.left_indent = Inches(0.2)
            continue

        # GFM Table
        if (stripped.startswith('|') and i + 1 < len(lines)
                and re.match(r'^\|[\s\-:|]+\|\s*$', lines[i + 1].strip())):
            header = [c.strip() for c in stripped.strip().strip('|').split('|')]
            i += 2
            rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                cells = [c.strip() for c in lines[i].strip().strip('|').split('|')]
                rows.append(cells)
                i += 1
            n = len(header)
            table = doc.add_table(rows=1 + len(rows), cols=n)
            table.style = 'Light Grid Accent 1'
            for j, h_text in enumerate(header):
                cell = table.cell(0, j)
                cell.text = ""
                _docx_fill_runs(cell.paragraphs[0], h_text, bold=True)
            for r_idx, row in enumerate(rows):
                row = (row + [''] * n)[:n]
                for j, cell_text in enumerate(row):
                    cell = table.cell(r_idx + 1, j)
                    cell.text = ""
                    _docx_fill_runs(cell.paragraphs[0], cell_text)
            continue

        # Horizontal rule
        if re.match(r'^(?:---+|\*\*\*+|___+)\s*$', stripped):
            p = doc.add_paragraph('_' * 60)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            i += 1
            continue

        # Bullet list
        bullet = re.match(r'^[-*+]\s+\[([ xX])\]\s+(.*)$', stripped)
        if bullet:
            mark = "☒" if bullet.group(1).lower() == 'x' else "☐"
            p = doc.add_paragraph()
            p.add_run(f"{mark} ")
            _docx_fill_runs(p, bullet.group(2))
            i += 1
            continue

        b = re.match(r'^[-*+]\s+(.*)$', stripped)
        if b:
            p = doc.add_paragraph(style='List Bullet')
            _docx_fill_runs(p, b.group(1))
            i += 1
            continue

        # Ordered list
        o = re.match(r'^\d+\.\s+(.*)$', stripped)
        if o:
            p = doc.add_paragraph(style='List Number')
            _docx_fill_runs(p, o.group(1))
            i += 1
            continue

        # Block quote
        if stripped.startswith('>'):
            buf = []
            while i < len(lines) and lines[i].strip().startswith('>'):
                quote_text = lines[i].strip()[1:]
                if quote_text.startswith(' '):
                    quote_text = quote_text[1:]
                if quote_text:
                    buf.append(quote_text)
                i += 1
            if buf:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.4)
                _docx_fill_runs(p, ' '.join(buf))
                for run in p.runs:
                    run.italic = True
            continue

        # Image
        img = re.match(r'^!\[([^\]]*)\]\((?:<([^>]*)>|([^)\s]+))\)\s*$', stripped)
        if img:
            url = img.group(2) or img.group(3)
            try:
                if os.path.exists(url):
                    doc.add_picture(url, width=Inches(5))
                else:
                    doc.add_paragraph(t("[Ảnh: {url}]", url=url))
            except Exception:
                doc.add_paragraph(t("[Ảnh: {url}]", url=url))
            if img.group(1):
                cap = doc.add_paragraph(img.group(1))
                cap.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                if cap.runs:
                    cap.runs[0].italic = True
            i += 1
            continue

        # Dòng trống
        if not stripped:
            i += 1
            continue

        # Paragraph thường
        buf = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt:
                break
            if (re.match(r'^(#{1,6})\s', nxt) or re.match(r'^[-*+]\s', nxt)
                    or re.match(r'^\d+\.\s', nxt) or nxt.startswith('>')
                    or nxt.startswith('|') or nxt.startswith('```')):
                break
            buf.append(nxt)
            i += 1
        p = doc.add_paragraph()
        _docx_fill_runs(p, ' '.join(buf))

    doc.save(path)
    return path


def _docx_strip_inline(text):
    """Bỏ ký hiệu markdown đơn giản trong heading/list text (giữ tiếng)."""
    text = re.sub(r'\$\$([\s\S]+?)\$\$', r'\1', text)
    text = re.sub(r'(?<!\$)\$([^\n$]+?)\$(?!\$)', r'\1', text)
    text = re.sub(r'\*\*([^\*]+)\*\*', r'\1', text)
    text = re.sub(r'__([^_]+)__', r'\1', text)
    text = re.sub(r'\*([^\*]+)\*', r'\1', text)
    text = re.sub(r'(?<!_)_([^_]+)_(?!_)', r'\1', text)
    text = re.sub(r'`([^`]+)`', r'\1', text)
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    return text


def _docx_fill_runs(paragraph, text, bold=False):
    """Thêm text vào paragraph, hỗ trợ inline markdown và $math$ bằng OMML."""
    if not text:
        return

    def apply_base(run):
        if bold:
            run.bold = True
        return run

    # Tokenize theo các marker phổ biến
    pattern = re.compile(
        r"((?<!\$)\$[^\n$]+?\$(?!\$)|"
        r"\*\*[^\*\n]+?\*\*|__[^_\n]+?__|"
        r"(?<![\*])\*[^\*\n]+?\*(?!\*)|(?<![_a-zA-Z])_[^_\n]+?_(?![_a-zA-Z])|"
        r"`[^`\n]+?`|\[[^\]]+?\]\((?:<[^>]+?>|[^)\s]+?)\))"
    )
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            apply_base(run)
        elif part.startswith('$') and part.endswith('$'):
            _docx_append_math(paragraph, part[1:-1])
        elif part.startswith('__') and part.endswith('__'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            apply_base(run)
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            apply_base(run)
        elif part.startswith('_') and part.endswith('_') and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            apply_base(run)
        elif part.startswith('`') and part.endswith('`'):
            from docx.shared import Pt
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            apply_base(run)
        elif part.startswith('[') and ']' in part and part.endswith(')'):
            m = re.match(r'\[([^\]]+)\]\((?:<([^>]*)>|([^)\s]+))\)', part)
            if m:
                run = paragraph.add_run(m.group(1))
                run.font.color.rgb = _docx_link_color()
                run.underline = True
                apply_base(run)
            else:
                apply_base(paragraph.add_run(part))
        else:
            apply_base(paragraph.add_run(part))


def _docx_link_color():
    from docx.shared import RGBColor
    return RGBColor(0x25, 0x63, 0xEB)


_OMML_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"


def _docx_append_math(paragraph, latex):
    """Append a Word Office Math (OMML) object to a python-docx paragraph."""
    latex = (latex or "").strip()
    if not latex:
        return
    try:
        from docx.oxml import parse_xml
        paragraph._p.append(parse_xml(_latex_to_omml(latex)))
    except Exception:
        # Fallback keeps exported content readable if XML insertion ever fails.
        paragraph.add_run(latex)


def _latex_to_omml(latex):
    parser = _OmmlLatexParser(latex)
    body = ''.join(parser.parse())
    if not body:
        body = _omml_run(latex)
    return f'<m:oMath xmlns:m="{_OMML_NS}">{body}</m:oMath>'


def _omml_text(text):
    return html.escape(text or "", quote=False)


def _omml_run(text):
    if not text:
        return ""
    return f'<m:r><m:t xml:space="preserve">{_omml_text(text)}</m:t></m:r>'


def _omml_wrap(tag, nodes):
    return f'<m:{tag}>{"".join(nodes)}</m:{tag}>'


def _omml_script(base, sub=None, sup=None):
    base_xml = _omml_wrap('e', base)
    if sub is not None and sup is not None:
        return (
            f'<m:sSubSup>{base_xml}'
            f'{_omml_wrap("sub", sub)}{_omml_wrap("sup", sup)}'
            f'</m:sSubSup>'
        )
    if sub is not None:
        return f'<m:sSub>{base_xml}{_omml_wrap("sub", sub)}</m:sSub>'
    if sup is not None:
        return f'<m:sSup>{base_xml}{_omml_wrap("sup", sup)}</m:sSup>'
    return ''.join(base)


def _omml_fraction(num, den):
    return (
        '<m:f><m:fPr><m:type m:val="bar"/></m:fPr>'
        f'{_omml_wrap("num", num)}{_omml_wrap("den", den)}'
        '</m:f>'
    )


def _omml_radical(expr, degree=None):
    if degree:
        return (
            '<m:rad><m:radPr><m:degHide m:val="0"/></m:radPr>'
            f'{_omml_wrap("deg", degree)}{_omml_wrap("e", expr)}'
            '</m:rad>'
        )
    return (
        '<m:rad><m:radPr><m:degHide m:val="1"/></m:radPr>'
        '<m:deg/>'
        f'{_omml_wrap("e", expr)}'
        '</m:rad>'
    )


_LATEX_SYMBOLS = {
    "alpha": "α", "beta": "β", "gamma": "γ", "delta": "δ",
    "epsilon": "ϵ", "varepsilon": "ε", "zeta": "ζ", "eta": "η",
    "theta": "θ", "vartheta": "ϑ", "iota": "ι", "kappa": "κ",
    "lambda": "λ", "mu": "μ", "nu": "ν", "xi": "ξ", "pi": "π",
    "rho": "ρ", "sigma": "σ", "tau": "τ", "upsilon": "υ",
    "phi": "ϕ", "varphi": "φ", "chi": "χ", "psi": "ψ", "omega": "ω",
    "Gamma": "Γ", "Delta": "Δ", "Theta": "Θ", "Lambda": "Λ",
    "Xi": "Ξ", "Pi": "Π", "Sigma": "Σ", "Upsilon": "Υ",
    "Phi": "Φ", "Psi": "Ψ", "Omega": "Ω",
    "times": "×", "cdot": "·", "div": "÷", "pm": "±", "mp": "∓",
    "le": "≤", "leq": "≤", "ge": "≥", "geq": "≥", "neq": "≠",
    "approx": "≈", "sim": "∼", "equiv": "≡", "infty": "∞",
    "sum": "∑", "prod": "∏", "int": "∫", "oint": "∮",
    "partial": "∂", "nabla": "∇", "forall": "∀", "exists": "∃",
    "in": "∈", "notin": "∉", "subset": "⊂", "subseteq": "⊆",
    "cup": "∪", "cap": "∩", "to": "→", "rightarrow": "→",
    "leftarrow": "←", "Rightarrow": "⇒", "Leftarrow": "⇐",
    "leftrightarrow": "↔", "sqrt": "√",
}


class _OmmlLatexParser:
    """Small LaTeX math parser for the Markdown DOCX exporter."""

    def __init__(self, text):
        self.text = text or ""
        self.pos = 0

    def parse(self):
        return self._parse_expression()

    def _parse_expression(self, stop_char=None):
        nodes = []
        while self.pos < len(self.text):
            ch = self.text[self.pos]
            if stop_char and ch == stop_char:
                self.pos += 1
                break
            if ch == '}':
                if stop_char:
                    self.pos += 1
                break
            atom = self._parse_atom(stop_char=stop_char)
            if atom:
                nodes.extend(self._apply_scripts(atom))
        return nodes

    def _parse_atom(self, stop_char=None):
        if self.pos >= len(self.text):
            return []
        ch = self.text[self.pos]
        if ch == '{':
            self.pos += 1
            return self._parse_expression(stop_char='}')
        if ch == '\\':
            return self._parse_command()
        if ch in '^_':
            self.pos += 1
            return [_omml_run(ch)]
        buf = []
        while self.pos < len(self.text):
            ch = self.text[self.pos]
            if ch == stop_char or ch in '{}\\^_':
                break
            buf.append(ch)
            self.pos += 1
        return [_omml_run(''.join(buf))]

    def _parse_command(self):
        self.pos += 1
        if self.pos >= len(self.text):
            return [_omml_run('\\')]

        if not self.text[self.pos].isalpha():
            ch = self.text[self.pos]
            self.pos += 1
            if ch in "{}_^$#&%":
                return [_omml_run(ch)]
            if ch in ",;:!":
                return [_omml_run(" ")]
            return [_omml_run(ch)]

        start = self.pos
        while self.pos < len(self.text) and self.text[self.pos].isalpha():
            self.pos += 1
        cmd = self.text[start:self.pos]

        if cmd in ("left", "right"):
            return []
        if cmd in ("quad", "qquad"):
            return [_omml_run("  " if cmd == "quad" else "    ")]
        if cmd == "frac":
            num = self._required_arg()
            den = self._required_arg()
            return [_omml_fraction(num, den)]
        if cmd == "sqrt":
            degree = self._optional_bracket_arg()
            expr = self._required_arg()
            return [_omml_radical(expr, degree)]
        if cmd in ("text", "mathrm", "mathbf", "mathit"):
            return [_omml_run(self._read_group_text())]
        if cmd in ("begin", "end"):
            self._read_group_text()
            return []

        return [_omml_run(_LATEX_SYMBOLS.get(cmd, "\\" + cmd))]

    def _apply_scripts(self, base):
        sub = None
        sup = None
        while self.pos < len(self.text) and self.text[self.pos] in ('_', '^'):
            marker = self.text[self.pos]
            self.pos += 1
            value = self._script_arg()
            if marker == '_':
                sub = value
            else:
                sup = value
        if sub is not None or sup is not None:
            return [_omml_script(base, sub=sub, sup=sup)]
        return base

    def _script_arg(self):
        self._skip_spaces()
        if self.pos < len(self.text) and self.text[self.pos] == '{':
            self.pos += 1
            return self._parse_expression(stop_char='}')
        if self.pos < len(self.text) and self.text[self.pos] == '\\':
            return self._parse_command()
        if self.pos < len(self.text):
            ch = self.text[self.pos]
            self.pos += 1
            return [_omml_run(ch)]
        return self._parse_atom()

    def _required_arg(self):
        self._skip_spaces()
        if self.pos < len(self.text) and self.text[self.pos] == '{':
            self.pos += 1
            return self._parse_expression(stop_char='}')
        return self._parse_atom()

    def _optional_bracket_arg(self):
        self._skip_spaces()
        if self.pos >= len(self.text) or self.text[self.pos] != '[':
            return None
        self.pos += 1
        start = self.pos
        depth = 1
        while self.pos < len(self.text) and depth:
            if self.text[self.pos] == '[':
                depth += 1
            elif self.text[self.pos] == ']':
                depth -= 1
                if depth == 0:
                    content = self.text[start:self.pos]
                    self.pos += 1
                    return _OmmlLatexParser(content).parse()
            self.pos += 1
        return None

    def _read_group_text(self):
        self._skip_spaces()
        if self.pos >= len(self.text) or self.text[self.pos] != '{':
            return ""
        self.pos += 1
        depth = 1
        out = []
        while self.pos < len(self.text) and depth:
            ch = self.text[self.pos]
            if ch == '{':
                depth += 1
                out.append(ch)
            elif ch == '}':
                depth -= 1
                if depth:
                    out.append(ch)
            else:
                out.append(ch)
            self.pos += 1
        return ''.join(out)

    def _skip_spaces(self):
        while self.pos < len(self.text) and self.text[self.pos].isspace():
            self.pos += 1
